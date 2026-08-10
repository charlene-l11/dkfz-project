from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


SOURCE = Path("DKFZ_median_spectrum_methodological_manuscript.docx")
OUTPUT = Path("DKFZ_median_spectrum_methodological_manuscript_revised.docx")


def replace_text(paragraph, text):
    """Replace paragraph contents while retaining paragraph-level formatting."""
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


doc = Document(SOURCE)

replacements = {
    6: "Draft derived from the supplied merged DKFZ pipeline and immutable HTC source tree • 7 August 2026",
    9: (
        "Hyperspectral imaging produces dense spectral measurements that may support objective tissue recognition, but methodological comparisons are easily confounded by information leakage, inconsistent preprocessing, changing model capacity, and incomplete provenance. We describe a configuration-driven pipeline for classifying organ-specific median spectra with the Hyperspectral Tissue Classification (HTC) median-pixel network. The software discovers precomputed median-spectrum comma-separated value files linked to configured HyperGUI and labelling artifacts, validates a common wavelength grid, and assigns complete subjects—not individual spectra—to training, validation, and testing partitions. The active reference configuration targets six rat-organ classes, 100 channels spanning 500–995 nm, and a 60%/20%/20% subject-level split. It additionally defines an external nine-class pig-organ test set. Selective-band experiments preserve a constant 100-channel model input: selected channels are standardized using training-only statistics and scattered back to their original positions, whereas unselected channels are zero-masked. The classifier is a one-dimensional convolutional network with three convolution–batch-normalization–exponential-linear-unit–average-pooling blocks, two fully connected layers, dropout, and a multiclass output head. Training uses Adam, exponential learning-rate decay, cross-entropy loss, deterministic seeding, validation-based checkpoint selection, and balanced oversampling. The pipeline exports manifests, split workbooks, predictions, raw and row-normalized confusion matrices, accuracy, balanced accuracy, macro recall, macro F1, class-specific precision, recall and F1, worst-class recall, software versions, source hashes, and fully resolved parameters. This design provides an auditable basis for wavelength-ablation studies and median-spectrum organ classification while explicitly preventing subject leakage and model-capacity drift."
    ),
    14: (
        "This manuscript describes the implemented workflow, specifies the reference experiment encoded in the repository, and defines a reporting framework for subsequent empirical results. Because the supplied repository contains no completed scenario outputs or clinical acquisition protocol, performance values, cohort sizes, acquisition details, run time and ethics statements cannot be inferred and remain to be inserted from the study record."
    ),
    17: (
        "The pipeline implements supervised multiclass classification of organ-specific median reflectance spectra. Each row of an internal manifest corresponds to one spectrum file and includes the file path, class label, integer label index, subject identifier, acquisition timestamp, image identifier and annotation name. The active internal reference configuration defines six rat-organ classes—stomach, small bowel, colon, liver, pancreas and kidney—and the external evaluation configuration defines nine pig-organ classes—stomach, small bowel, colon, liver, pancreas, kidney, spleen, bladder and omentum. External labels define evaluation-matrix rows and are not supplied to the trained model loss. The biological source, number of subjects and acquisitions, inclusion and exclusion criteria, and reference-standard annotation procedure must be supplied from the experimental protocol before submission."
    ),
    28: (
        "Subjects, rather than spectra, are the indivisible units of partitioning. The configured target fractions are 0.60 for training, 0.20 for validation and 0.20 for testing. Integer partition sizes are obtained by rounding the validation and test targets, with at least one subject assigned to each, and assigning the remainder to training. Before assignment, each class must occur in at least as many distinct subjects as the number of partitions required to contain all classes. Using split seed 42, the algorithm evaluates as many as 20,000 random subject permutations. Candidates are ranked first by the number of missing class–partition combinations and subsequently by the sum, over partitions, of the coefficient of variation of class-specific spectrum counts. The lowest-scoring feasible assignment is retained. No subject can occur in more than one partition, and all six internal classes are required in training, validation and testing. When external testing is enabled, the external cohort replaces the internal test manifest; it remains independent of model fitting and may contain a different class set."
    ),
    33: (
        "The classifier reuses the HTC ModelPixel architecture with a single spectral input channel. Three valid one-dimensional convolutions use kernel size 5 and output 64, 32 and 16 feature maps, respectively. Each convolution is followed by one-dimensional batch normalization, an exponential linear unit activation and average pooling with kernel size 2. The convolutional representation is flattened and adaptively reduced to the dimension determined for the configured input width. Two dense layers contain 100 and 50 units, each followed by batch normalization, exponential linear unit activation and dropout with probability 0.2. A task-specific classification head maps the 50-dimensional representation to six logits in the active internal experiment. External evaluation retains these six output columns even when the external matrix contains nine true-class rows. This architecture is inherited directly from the authoritative HTC runtime."
    ),
    35: (
        "The integration layer does not implement the gpu_smoke, gpu_practical or gpu_paper profiles shown in the preliminary notes; instead, all reported settings are read directly from the active YAML configuration and saved again in the resolved run parameters. Training proceeds for at most 150 epochs using Adam with learning rate 2×10⁻⁴ and zero weight decay, followed by an ExponentialLR scheduler with multiplicative factor 0.9. The batch size is 64, numerical precision is 32-bit (32-true), and one device is selected with the accelerator set to auto. The number of data-loader workers is zero. Randomness is initialized with seed 42, including worker initialization. The training epoch size is not a fixed numerical constant in the YAML; it resolves at run time to the number of training spectra. With balanced oversampling, this number of observations is sampled with replacement in each epoch. The objective is multiclass cross-entropy loss."
    ),
    38: (
        "At every epoch, HTC aggregates validation confusion matrices by subject and logs mean subject-level accuracy. The checkpoint attaining maximum validation accuracy is retained, together with the most recent checkpoint. After fitting, the best checkpoint is used to generate row-level validation predictions and to evaluate the held-out test loader. Class predictions are obtained as the argmax of the model logits. For external evaluation, prediction is performed without calculating an HTC test loss because external class identifiers need not correspond one-to-one to the training targets."
    ),
    39: (
        "For both validation and test partitions, the pipeline calculates a count confusion matrix C and its row-normalized form, Cᵢⱼ/ΣⱼCᵢⱼ. Matrix rows denote true classes and columns denote trained-model classes; consequently, the active external experiment yields a 9 × 6 matrix. Let N be the number of evaluated spectra, TPᵢ = Cᵢᵢ, nᵢ = ΣⱼCᵢⱼ and pᵢ = ΣⱼCⱼᵢ for a class name present in both axes. Overall accuracy is ΣᵢTPᵢ/N. Class-specific precision is TPᵢ/pᵢ, recall (sensitivity) is TPᵢ/nᵢ and F1 is 2·precisionᵢ·recallᵢ/(precisionᵢ+recallᵢ); a zero denominator yields zero. For an external-only class with no matching model-output column, TPᵢ, precision, recall and F1 are defined as zero. Balanced accuracy and macro recall are identical in this implementation and equal the unweighted mean class recall, whereas macro F1 is the unweighted mean class F1. The output further records N, the number of true classes, the number of predicted classes, the class with minimum recall and that worst-class recall; alphabetical order breaks a tie in minimum recall. Per-class output additionally reports label index, support nᵢ, correct count TPᵢ, the most frequently confused predicted class and its count. All metrics are spectrum-level because each manifest row is one median spectrum. Specificity, negative predictive value, micro- or weighted-average scores, receiver-operating-characteristic area under the curve, calibration indices, confidence intervals and subject-level aggregate metrics are not calculated by the current code and therefore are not reported as implemented outcomes."
    ),
    41: (
        "Optional forward and reverse analyses quantify how performance changes as contiguous spectral coverage expands. Forward runs begin with the first two eligible bands at the lower boundary and add one channel per run until the configured forward stop; reverse runs begin with the last two eligible bands and expand downward until the reverse stop. Both modes are disabled in the active reference configuration, whose stop values are 995 and 500 nm, respectively. Discovery and subject splitting occur once at the master level and are reused by every child run. Each child receives a generated YAML configuration and a fixed-width selected-band representation. Completed children are recognized from their run-parameter record and skipped on restart. Aggregate tables contain validation and test accuracy together with selected-channel count and boundary wavelength, and the corresponding accuracy trajectories are plotted."
    ),
    47: (
        "The current code package contains the analysis implementation but no completed scenario outputs; accordingly, no empirical performance estimate can be stated from the repository alone. A completed report should reproduce all metrics written by the software: validation and test accuracy, balanced accuracy (macro recall), macro F1, evaluated-spectrum count, true- and predicted-class counts, worst class and worst-class recall, as well as class-specific support, correct count, precision, recall, F1 and dominant confusion. It should also report subjects and spectra by class and partition, the selected wavelength grid, training duration and hardware, the validation-selected epoch, and raw and row-normalized confusion matrices. Confidence intervals or subject-level distributions require an additional prespecified analysis because the present implementation does not compute them. Wavelength experiments should present validation and test trajectories separately and identify any selection made after inspecting test performance as exploratory."
    ),
}

for index, text in replacements.items():
    replace_text(doc.paragraphs[index], text)

# Correct and expand the reference-settings table.
table = doc.tables[0]
values = [
    ("Spectral input", "100 channels, 500–995 nm; all bands selected", "Common grid for discovery and wavelength variants"),
    ("Internal classes", "6 rat-organ labels", "Stomach, small bowel, colon, liver, pancreas and kidney"),
    ("External classes", "9 pig-organ labels", "Independent evaluation; 9 true-class rows × 6 output columns"),
    ("Partitioning", "60% train / 20% validation / 20% test", "Subject-level; seed 42; ≤20,000 candidate assignments"),
    ("Training", "≤150 epochs; batch 64; learning rate 2×10⁻⁴", "One device; accelerator auto; 32-true precision; 0 workers"),
    ("Epoch size", "Number of training spectra at run time", "Replacement sampling under balanced oversampling"),
    ("Normalization", "Per-channel z-standardization", "Mean and population SD estimated from training spectra only"),
    ("Imbalance", "Balanced oversampling", "Mutually exclusive with class-weighted loss"),
    ("Checkpoint", "Maximum validation accuracy", "Best checkpoint used for validation prediction and final testing"),
    ("Summary metrics", "Accuracy; balanced accuracy/macro recall; macro F1", "Plus row count, class counts, worst class and worst-class recall"),
    ("Class metrics", "Support; correct; precision; recall; F1", "Plus dominant confused class and confusion count"),
    ("Outputs", "CSV, Excel, PNG, PDF and JSON; 300 dpi", "Predictions, matrices, metrics, splits and provenance"),
]
while len(table.rows) < len(values) + 1:
    table.add_row()
while len(table.rows) > len(values) + 1:
    table._tbl.remove(table.rows[-1]._tr)
for ri, row_values in enumerate(values, start=1):
    for ci, value in enumerate(row_values):
        table.cell(ri, ci).text = value

# Avoid splitting compact table rows across pages and repeat the header.
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if ri == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)

doc.save(OUTPUT)
print(OUTPUT.resolve())
